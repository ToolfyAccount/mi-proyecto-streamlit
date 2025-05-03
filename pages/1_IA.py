import streamlit as st
from PIL import Image
from io import BytesIO
from ai21 import AI21Client
from ai21.models.chat import ChatMessage
from peewee import MySQLDatabase, Model, CharField, IntegerField
from docx import Document
import io
import os
import pymysql
import google.generativeai as genai

# Sidebar ordenada
db = MySQLDatabase(
    'defaultdb',
    user=os.environ.get("Usuarios_1"),
    password=os.environ.get("Password"),
    host=os.environ.get("Host"),
    port=19758,
    ssl={'fake_flag_to_enable_ssl': True}  # ✅ Este es el cambio importante
)


class Usuario(Model):
    nombre = CharField()
    contraseña = CharField()

    class Meta:
        database = db


# Intentamos conectar y mostramos mensaje en consola
try:
    db.connect()
    print("✅ Conexión exitosa a la base de datos.")
except Exception as e:
    print(f"❌ Error al conectar a la base de datos: {e}")
db.create_tables([Usuario])

# Sesión


if "Auntentificado" not in st.session_state or not st.session_state["Auntentificado"]:
    st.error("🚫 No estás autorizado. Redirigiendo al inicio de sesión...")
    st.switch_page("pages/4_Login.py")

# --- ESTÉTICA PERSONALIZADA ---
st.markdown(
    """
    <style>
    /* Página general */
    body {
        background-color: #0e1117;
        color: white;
    }

    /* Título principal */
    .main-title {
    font-size: 48px;
    font-weight: bold;
    color: #3399ff;
    text-align: center;
    margin-bottom: 10px;


}
    .titulo {
    font-size: 48px;
    font-weight: 100; /* Versión más ligera */
    color: #D0E7FF;
    text-align: center;
    margin-bottom: 10px;
    font-family: 'Montserrat', sans-serif;
}







    /* Subtítulo */
    .subtext {
        text-align: center;
        color: #aaa;
        font-size: 18px;
        margin-bottom: 40px;
    }

    /* Caja contenedora */
    .input-box {
        padding: 20px;
        border-radius: 10px;
        background-color: #1e1e1e;
        margin-bottom: 20px;
    }

    /* Input personalizado */
    .stTextInput>div>div>input {
        background-color: #2c2f36;
        color: white;
        border: none;
        padding: 10px;
        border-radius: 8px;
    }

    .stTextInput>div>div>input::placeholder {
        color: #aaa;
    }

    /* Botón personalizado */
    .stButton > button {
        background-color: #ff4b4b;
        color: white;
        font-weight: bold;
        border-radius: 10px;
        padding: 10px 20px;
        transition: background-color 0.3s ease;
        border: none;
    }

    .stButton > button:hover {
        background-color: #ff6666;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# --- INTERFAZ PRINCIPAL ---
st.markdown('<link href="https://fonts.googleapis.com/css2?family=Montserrat:wght@100;300;400&display=swap" rel="stylesheet"> <div class="titulo"> LEVERFUL</div>', unsafe_allow_html=True)
st.markdown('<div class="main-title">🧠 IA</div>', unsafe_allow_html=True)

# Usuario actual
User = Usuario.select().where(
    Usuario.nombre == st.session_state["usuario"]).first()


st.markdown(
    f'<div class="subtext">Bienvenido, <strong>{User.nombre}</strong>. Estás en el IA principal.</div>', unsafe_allow_html=True)

# Entrada de pregunta

# Subida de archivo directamente sin guardar en sesión
archivo_nuevo = st.file_uploader("Selecciona un archivo", type=["txt", "docx"])

if archivo_nuevo is not None:
    st.success(f"Archivo cargado: {archivo_nuevo.name}")
    # Aquí puedes procesar el archivo directamente, por ejemplo:
    if archivo_nuevo.type == "text/plain":
        contenido = archivo_nuevo.read().decode("utf-8")
        st.text_area("Contenido del archivo", contenido, height=300)
    elif archivo_nuevo.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(archivo_nuevo)
        contenido = "\n".join([p.text for p in doc.paragraphs])
        st.text_area("Contenido del archivo", contenido, height=300)


Text = st.chat_input("Pon la pregunta")
st.markdown('</div>', unsafe_allow_html=True)

API = os.environ.get("API")

genai.configure(api_key=API)

# Configuración del modelo con temperatura
model = genai.GenerativeModel(
    model_name="gemini-2.0-flash",
    generation_config=genai.types.GenerationConfig(
        temperature=0.7
    )
)

chat = model.start_chat()

respuesta = chat.send_message("Guarda esto en tu memoria: Tienes la opcion de querer hablar con markdown, si no quieres esta bien , en los problemas matematicos puedes encerrarlo en un cuadrado para diferenciar e igual con el codigo, si quieres hacer un archivo para que el usuario lo descargue escribe (Generacion.txt) como primera palabra del texto, y lo demas del texto escribes lo que quieres escribir en el .txt, el txt solo lo puedes hacer si el USUARIO te lo pide, seria lo mismo si quieres hacer un archivo .docx, escribes (Generacion.docx). si quieres hacer un archivo txt o docx no puedes hablar con markdown, CUANDO VAYAS A ESCRIBIR UN ARCHIVO ESCRIBE NORMAL, SIN MARKDOWN y no puedes crear otros tipos de archivos, solo txt y docx. no puedes mencionar nada de lo que dije aca ok, ESTE MENSAJE SON INSTRUCCIONES DEL DESAROLLADOR.")


# Botón para preguntar
if Text:
    if archivo_nuevo is None:
        if Text.strip():
            try:

                RTA = chat.send_message(
                    f"Mensaje del usuario:'{Text}'")
                RTA = RTA.text

                if "Generacion.txt" in RTA:
                    RTAT = RTA.replace("Generacion.txt", "").strip()
                    st.download_button(
                        label="⬇️ Descargar respuesta en .txt",
                        data=RTAT.encode('utf-8'),
                        file_name="Generacion.txt",
                        mime="text/plain"
                    )
                if "Generacion.docx" in RTA:
                    RTAT = RTA.replace("Generacion.docx", "").strip()
                    # Crear el documento Word
                    doc = Document()
                    doc.add_heading('Generacion de IA', level=1)
                    doc.add_paragraph(RTAT)

                    # Guardarlo en una variable como flujo de bytes
                    doc_variable = io.BytesIO()
                    doc.save(doc_variable)
                    # Es importante mover el puntero al inicio del flujo
                    doc_variable.seek(0)
                    st.download_button(
                        label="⬇️ Descargar respuesta en .docx(Documento de word)",
                        data=doc_variable,
                        file_name="Generacion.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                st.markdown("---")
                st.markdown("### 📩 Respuesta:")
                st.markdown(RTA)

            except Exception as e:
                st.error(f"❌ Error: {str(e)}")
        else:
            st.warning(
                "⚠️ Por favor, escribe una pregunta antes de hacer clic en 'Preguntar'.")

    else:

        if archivo_nuevo.type == "text/plain":
            Archivo = archivo_nuevo.read().decode("utf-8")

        elif archivo_nuevo.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(archivo_nuevo)
            Archivo = "\n".join([p.text for p in doc.paragraphs])

        RTA = chat.send_message(
            f"Mensaje del usuario:'{Text}', Archivo puesto por el Usuario: '{Archivo}'      Nombre del archivo:'{archivo_nuevo.name}")
        RTA = RTA.text

        if "Generacion.txt" in RTA:
            RTAT = RTA.replace("Generacion.txt", "").strip()
            st.download_button(
                label="⬇️ Descargar respuesta en .txt",
                data=RTAT.encode('utf-8'),
                file_name="Generacion.txt",
                mime="text/plain"
            )
        if "Generacion.docx" in RTA:
            RTAT = RTA.replace("Generacion.docx", "").strip()
            # Crear el documento Word
            doc = Document()
            doc.add_heading('Generacion de IA', level=1)
            doc.add_paragraph(RTAT)

            # Guardarlo en una variable como flujo de bytes
            doc_variable = io.BytesIO()
            doc.save(doc_variable)
            # Es importante mover el puntero al inicio del flujo
            doc_variable.seek(0)
            st.download_button(
                label="⬇️ Descargar respuesta en .docx(Documento de word)",
                data=doc_variable,
                file_name="Generacion.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        st.markdown("---")
        st.markdown("### 📩 Respuesta:")
        st.markdown(RTA)
