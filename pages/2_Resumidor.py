import streamlit as st
from PIL import Image
from io import BytesIO
from ai21 import AI21Client
from ai21.models.chat import ChatMessage
from peewee import MySQLDatabase, Model, CharField, IntegerField
from docx import Document
import io
import os

def texto_despues_del_punto(texto):
    # Encontrar la posición del primer punto
    punto_pos = texto.find('.')
    
    # Si hay un punto, devolver el texto después del primer punto
    if punto_pos != -1:
        return texto[punto_pos + 1:].strip()  # Eliminar espacios extra
    else:
        return "No hay punto en el texto."
# Configuración de base de datos
db = MySQLDatabase(
    'defaultdb',
    user=os.environ.get("Usuarios_1"),
    password=os.environ.get("Password"),
    host=os.environ.get("Host"),
    port=19758,
    ssl={'fake_flag_to_enable_ssl': True}  # ✅ Este es el cambio importante
)
class Usuario(Model):
    nombre = CharField()
    contraseña = CharField()
    Api = CharField()

    class Meta:
        database = db

db.connect()
db.create_tables([Usuario])



if "Auntentificado" not in st.session_state or not st.session_state["Auntentificado"]:
    st.error("🚫 No estás autorizado. Redirigiendo al inicio de sesión...")
    st.switch_page("pages/4_Login.py")
# --- ESTÉTICA PERSONALIZADA ---
st.markdown(
    """
    <style>
    /* Página general */
    body {
        background-color: #0e1117;
        color: white;
    }

    /* Título principal */
    .main-title {
    font-size: 48px;
    font-weight: bold;
    color: #3399ff;
    text-align: center;
    margin-bottom: 10px;


}
    .titulo {
    font-size: 48px;
    font-weight: 100; /* Versión más ligera */
    color: #D0E7FF;
    text-align: center;
    margin-bottom: 10px;
    font-family: 'Montserrat', sans-serif;
}







    /* Subtítulo */
    .subtext {
        text-align: center;
        color: #aaa;
        font-size: 18px;
        margin-bottom: 40px;
    }

    /* Caja contenedora */
    .input-box {
        padding: 20px;
        border-radius: 10px;
        background-color: #1e1e1e;
        margin-bottom: 20px;
    }

    /* Input personalizado */
    .stTextInput>div>div>input {
        background-color: #2c2f36;
        color: white;
        border: none;
        padding: 10px;
        border-radius: 8px;
    }

    .stTextInput>div>div>input::placeholder {
        color: #aaa;
    }

    /* Botón personalizado */
    .stButton > button {
        background-color: #ff4b4b;
        color: white;
        font-weight: bold;
        border-radius: 10px;
        padding: 10px 20px;
        transition: background-color 0.3s ease;
        border: none;
    }

    .stButton > button:hover {
        background-color: #ff6666;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# --- INTERFAZ PRINCIPAL ---
st.markdown('<link href="https://fonts.googleapis.com/css2?family=Montserrat:wght@100;300;400&display=swap" rel="stylesheet"> <div class="titulo"> LEVERFUL</div>', unsafe_allow_html=True)
st.markdown('<div class="main-title">📚 Resumidor</div>', unsafe_allow_html=True)

# Usuario actual
User = Usuario.select().where(Usuario.nombre == st.session_state["usuario"]).first()
API = User.Api

st.markdown(f'<div class="subtext">Bienvenido, <strong>{User.nombre}</strong>. Estás en el Resumidor.</div>', unsafe_allow_html=True)

# Entrada de pregunta

# Subida de archivo directamente sin guardar en sesión
archivo_nuevo = st.file_uploader("Selecciona un archivo", type=["txt", "docx"])

if archivo_nuevo is not None:
    st.success(f"Archivo cargado: {archivo_nuevo.name}")
    # Aquí puedes procesar el archivo directamente, por ejemplo:
    if archivo_nuevo.type == "text/plain":
        contenido = archivo_nuevo.read().decode("utf-8")
        st.text_area("Contenido del archivo", contenido, height=300)
    elif archivo_nuevo.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(archivo_nuevo)
        contenido = "\n".join([p.text for p in doc.paragraphs])
        st.text_area("Contenido del archivo", contenido, height=300)

st.markdown("### ❓ Inserta tu texto a resumir (opcional si cargaste un archivo):")
Text = st.text_input(
    "Escribe aquí tu texto a resumir.",
    placeholder="Texto a resumir",
    disabled=archivo_nuevo is not None
)
st.markdown('</div>', unsafe_allow_html=True)

# Función para obtener respuesta
def Respuesta(mensajes):
    response = client.chat.completions.create(
        messages=mensajes,
        model="jamba-1.5-large",
        temperature=0,
        max_tokens=4090
    )
    return response.choices[0].message.content

# Botón para preguntar
if st.button("💬 Resumir.", type="primary"):
    if archivo_nuevo is not None:
        if archivo_nuevo.type == "text/plain":
            Text = archivo_nuevo.read().decode("utf-8")
                    
        elif archivo_nuevo.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(archivo_nuevo)
            Text = "\n".join([p.text for p in doc.paragraphs])
    if Text.strip():
        try:
            client = AI21Client(api_key=API)
            
            RTA = Respuesta([
                ChatMessage(role="user", content= (f"Resume el siguiente texto de forma clara y estructurada, conservando toda la información relevante, el contexto y los detalles esenciales. El resumen debe ser más corto que el texto original, pero no excesivamente breve. Asegúrate de incluir los puntos clave, hechos importantes, relaciones entre ideas y cualquier información crítica para comprender el contenido completo. Mensaje del usuario:{Text}"))
            ])
               
            
            
            st.download_button(
                label="⬇️ Descargar resumen en .txt",
                data=RTA.encode('utf-8'),
                file_name="Generacion.txt",
                mime="text/plain"
            )
            
                
                
            doc = Document()
            doc.add_heading('Generacion de IA', level=1)
            doc.add_paragraph(RTA)

            # Guardarlo en una variable como flujo de bytes
            doc_variable = io.BytesIO()
            doc.save(doc_variable)
            doc_variable.seek(0)  # Es importante mover el puntero al inicio del flujo
            st.download_button(
                label="⬇️ Descargar resumen en .docx",
                data=doc_variable,
                file_name="Generacion.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.markdown("---")
            st.markdown("### 📩 Respuesta:")
            st.markdown(RTA)

        except Exception as e:
            st.error(f"❌ Error: {str(e)}")
    else:
        st.warning("⚠️ Por favor, escribe una Resumen antes de hacer clic en 'Resumir'.")
